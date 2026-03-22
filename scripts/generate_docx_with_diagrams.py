#!/usr/bin/env python3
"""
VConfi Implementation Plan — DOCX Generator with Diagram Embedding

This enhanced version:
1. Converts Mermaid diagrams to PNG using mermaid-cli
2. Embeds the PNG images directly into the DOCX
3. Handles both inline base64 and external file references

Usage:
    python generate_docx_with_diagrams.py merge --parts Part*.md --output plan.docx
"""

import sys
import os
import re
import argparse
import base64
import tempfile
import subprocess
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx is required. Install: pip install python-docx")
    sys.exit(1)


# Brand colors
BRAND_PRIMARY = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_ACCENT = RGBColor(0x00, 0x7A, 0xCC)
BRAND_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BRAND_DARK = RGBColor(0x33, 0x33, 0x33)
TABLE_HEADER_BG = "007ACC"


def extract_and_render_diagrams(md_content, temp_dir):
    """
    Extract mermaid diagrams from markdown and render to PNG.
    
    Returns:
        (modified_content, diagram_files) where diagram_files is a dict
        mapping placeholder IDs to PNG file paths
    """
    pattern = r'```mermaid\n(.*?)```'
    diagrams = re.findall(pattern, md_content, re.DOTALL)
    
    if not diagrams:
        return md_content, {}
    
    diagram_files = {}
    modified_content = md_content
    
    for i, diagram in enumerate(diagrams):
        # Create temp mmd file
        mmd_path = os.path.join(temp_dir, f"diagram_{i}.mmd")
        png_path = os.path.join(temp_dir, f"diagram_{i}.png")
        
        with open(mmd_path, 'w') as f:
            f.write(diagram)
        
        # Render with mermaid-cli
        try:
            cmd = [
                'mmdc',
                '-i', mmd_path,
                '-o', png_path,
                '-w', '1200',
                '-h', '800',
                '-b', 'white'
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0 and os.path.exists(png_path):
                # Replace mermaid block with placeholder
                placeholder = f"[[DIAGRAM:{i}]]"
                old_block = f'```mermaid\n{diagram}```'
                modified_content = modified_content.replace(old_block, placeholder, 1)
                diagram_files[placeholder] = png_path
        except Exception as e:
            print(f"Warning: Failed to render diagram {i}: {e}")
    
    return modified_content, diagram_files


def add_image_to_doc(doc, image_path, width=Inches(6)):
    """Add an image to document with proper formatting."""
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        return True
    except Exception as e:
        print(f"Warning: Could not add image {image_path}: {e}")
        return False


def process_markdown_with_diagrams(md_content, doc, temp_dir):
    """Process markdown content, rendering and embedding diagrams."""
    # First, extract and render diagrams
    modified_content, diagram_files = extract_and_render_diagrams(md_content, temp_dir)
    
    # Split content by diagram placeholders
    parts = re.split(r'(\[\[DIAGRAM:\d+\]\])', modified_content)
    
    for part in parts:
        if part.startswith('[[DIAGRAM:') and part in diagram_files:
            # Insert image
            add_image_to_doc(doc, diagram_files[part])
        else:
            # Process as regular markdown
            parse_markdown_text(part, doc)


def parse_markdown_text(text, doc):
    """Parse markdown text (without diagrams) and add to document."""
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                doc.add_heading(text, level=level)
                i += 1
                continue
        
        # Tables
        if '|' in line and line.startswith('|'):
            headers, rows, i = parse_table(lines, i)
            if headers:
                add_table(doc, headers, rows)
            continue
        
        # Regular paragraph
        p = doc.add_paragraph(line)
        i += 1


def parse_table(lines, start_idx):
    """Parse a markdown table."""
    headers = []
    rows = []
    i = start_idx
    
    # Header row
    if i < len(lines) and '|' in lines[i]:
        parts = [p.strip() for p in lines[i].split('|')]
        headers = [p for p in parts if p]
        i += 1
    
    # Skip separator
    if i < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i]):
        i += 1
    
    # Data rows
    while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
        parts = [p.strip() for p in lines[i].split('|')]
        row = [p for p in parts if p != '']
        if row:
            rows.append(row)
        i += 1
    
    return headers, rows, i


def add_table(doc, headers, rows):
    """Add a formatted table to document."""
    if not headers or not rows:
        return
    
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        cell.text = header
    
    # Data
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data[:len(headers)]):
            table.cell(row_idx+1, col_idx).text = cell_text
    
    doc.add_paragraph()


def convert_md_to_docx(input_path, output_path, client_name=None, project_name=None):
    """Convert markdown to DOCX with diagram rendering."""
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    doc = Document()
    
    # Create temp directory for diagrams
    with tempfile.TemporaryDirectory() as temp_dir:
        # Process content with diagrams
        process_markdown_with_diagrams(md_content, doc, temp_dir)
    
    doc.save(output_path)
    print(f"Document saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='VConfi DOCX Generator with Diagrams')
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('output', nargs='?', help='Output DOCX file')
    parser.add_argument('--client', help='Client name')
    parser.add_argument('--project', help='Project name')
    
    args = parser.parse_args()
    
    if not args.output:
        args.output = os.path.splitext(args.input)[0] + '.docx'
    
    convert_md_to_docx(args.input, args.output, args.client, args.project)


if __name__ == '__main__':
    main()
