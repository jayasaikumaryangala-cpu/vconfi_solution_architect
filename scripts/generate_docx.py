#!/usr/bin/env python3
"""
VConfi Implementation Plan — DOCX Generator
Converts a markdown implementation plan into a professionally formatted Word document.

Usage:
    python generate_docx.py <input.md> [output.docx]
    python generate_docx.py <input.md> --client "Client Name" --project "Project Name"

If output is not specified, it replaces .md with .docx in the same directory.
"""

import sys
import os
import re
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


# ── Brand Colors ──────────────────────────────────────────────────────────────
BRAND_PRIMARY = RGBColor(0x1A, 0x1A, 0x2E)    # Dark navy
BRAND_ACCENT = RGBColor(0x00, 0x7A, 0xCC)      # Blue accent
BRAND_LIGHT = RGBColor(0xF0, 0xF4, 0xF8)       # Light background
BRAND_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BRAND_DARK = RGBColor(0x33, 0x33, 0x33)
BRAND_GRAY = RGBColor(0x66, 0x66, 0x66)
BRAND_RED = RGBColor(0xCC, 0x00, 0x00)
BRAND_GREEN = RGBColor(0x00, 0x80, 0x00)
BRAND_ORANGE = RGBColor(0xFF, 0x8C, 0x00)

TABLE_HEADER_BG = "007ACC"
TABLE_ALT_ROW = "F0F4F8"


def setup_styles(doc):
    """Configure document styles for professional output."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = BRAND_DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = BRAND_PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = BRAND_ACCENT
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = BRAND_PRIMARY
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.color.rgb = BRAND_ACCENT
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "007ACC")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_cover_page(doc, client_name, project_name, date_str):
    """Add a professional cover page."""
    # Add spacing before title
    for _ in range(6):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VConfi Solutions")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY
    run.font.name = 'Calibri'

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IT Infrastructure Solutions")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_GRAY
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = BRAND_ACCENT
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Implementation Plan")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BRAND_ACCENT
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Project name
    if project_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_name)
        run.font.size = Pt(18)
        run.font.color.rgb = BRAND_DARK
        run.font.name = 'Calibri'

    doc.add_paragraph()

    # Client and date info
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("Client", client_name or "[Client Name]"),
        ("Prepared By", "VConfi Solutions Team"),
        ("Date", date_str),
        ("Classification", "CONFIDENTIAL"),
    ]

    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.cell(i, 0)
        cell_value = info_table.cell(i, 1)

        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{label}:")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BRAND_GRAY
        run.font.name = 'Calibri'

        p = cell_value.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"  {value}")
        run.font.size = Pt(11)
        run.font.color.rgb = BRAND_DARK
        run.font.name = 'Calibri'
        if label == "Classification":
            run.font.color.rgb = BRAND_RED
            run.font.bold = True

    # Page break after cover
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add a Table of Contents page."""
    doc.add_heading("Table of Contents", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "This document contains the complete implementation plan. "
        "Refer to the section headings for navigation."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_GRAY
    run.font.italic = True

    # Add a TOC field (Word will populate it when opened)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run("Right-click and select 'Update Field' to populate the Table of Contents")
    run4.font.color.rgb = BRAND_GRAY
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    doc.add_page_break()


def add_formatted_table(doc, headers, rows):
    """Add a professionally styled table."""
    if not headers or not rows:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.strip())
        run.font.bold = True
        run.font.color.rgb = BRAND_WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx + 1, col_idx)
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = BRAND_DARK

    doc.add_paragraph()


def add_severity_badge(paragraph, severity):
    """Add a colored severity indicator."""
    severity_lower = severity.strip().lower()
    color_map = {
        'critical': BRAND_RED,
        'high': RGBColor(0xFF, 0x45, 0x00),
        'medium': BRAND_ORANGE,
        'low': BRAND_GREEN,
        'informational': BRAND_GRAY,
        'planned': BRAND_ACCENT,
    }
    color = color_map.get(severity_lower, BRAND_DARK)
    run = paragraph.add_run(f" [{severity.strip().upper()}]")
    run.font.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(9)


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = []
    rows = []
    i = start_idx

    # Parse header row
    if i < len(lines) and '|' in lines[i]:
        parts = [p.strip() for p in lines[i].split('|')]
        headers = [p for p in parts if p and not re.match(r'^[-:]+$', p)]
        i += 1

    # Skip separator row
    if i < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i]):
        i += 1

    # Parse data rows
    while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
        parts = [p.strip() for p in lines[i].split('|')]
        row = [p for p in parts if p != '']
        if row and not re.match(r'^[-:]+$', row[0]):
            rows.append(row)
        i += 1

    return headers, rows, i


def parse_and_convert(md_content, doc, client_name=None, project_name=None):
    """Parse markdown content and add to Word document."""
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    past_frontmatter = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = BRAND_DARK
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
                code_content = []
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Skip YAML frontmatter (--- delimited block at the start of the file)
        if stripped == '---' and not past_frontmatter:
            # Check if no real content has appeared before this line
            preceding_content = any(
                l.strip() and l.strip() != '---' for l in lines[:i]
            )
            if not preceding_content:
                i += 1
                while i < len(lines) and lines[i].strip() != '---':
                    i += 1
                i += 1
                past_frontmatter = True
                continue
        past_frontmatter = True

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                # Remove markdown bold markers
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                heading = doc.add_heading(text, level=level)
                i += 1
                continue

        # Tables
        if '|' in stripped and stripped.startswith('|'):
            headers, rows, end_idx = parse_markdown_table(lines, i)
            if headers and rows:
                add_formatted_table(doc, headers, rows)
            i = end_idx
            continue

        # Checkbox items
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            checked = stripped.startswith('- [x]')
            text = stripped[5:].strip()
            p = doc.add_paragraph(style='List Bullet')
            checkbox = "☑" if checked else "☐"
            run = p.add_run(f"{checkbox} ")
            run.font.size = Pt(11)
            run.font.color.rgb = BRAND_GREEN if checked else BRAND_GRAY
            # Handle bold text within
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                if idx % 2 == 1:
                    run.font.bold = True
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold text within bullets
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for idx, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = BRAND_DARK
                if idx % 2 == 1:
                    run.font.bold = True
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            text = num_match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for idx, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = BRAND_DARK
                if idx % 2 == 1:
                    run.font.bold = True
            i += 1
            continue

        # Blockquotes
        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for idx, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = BRAND_GRAY
                run.font.name = 'Calibri'
                if idx % 2 == 1:
                    run.font.bold = True
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        # Handle inline formatting (bold, italic, inline code)
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', stripped)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
                run.font.color.rgb = BRAND_DARK
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.font.italic = True
                run.font.color.rgb = BRAND_GRAY
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = BRAND_ACCENT
            else:
                run = p.add_run(part)
                run.font.color.rgb = BRAND_DARK

            run.font.size = Pt(10)
            if not run.font.name:
                run.font.name = 'Calibri'

        i += 1


def add_footer(doc):
    """Add footer with page numbers and company name."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run("VConfi Solutions  |  CONFIDENTIAL  |  Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = BRAND_GRAY
        run.font.name = 'Calibri'

        # Page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run2 = p.add_run()
        run2._r.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run3 = p.add_run()
        run3._r.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run4 = p.add_run()
        run4._r.append(fldChar2)


def set_page_margins(doc):
    """Set professional page margins."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_part_divider(doc, part_title):
    """Add a visual divider between parts."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = BRAND_ACCENT
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(part_title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = BRAND_ACCENT
    run.font.size = Pt(12)

    doc.add_paragraph()


def convert_md_to_docx(input_path, output_path=None, client_name=None, project_name=None):
    """Convert a single markdown file to docx."""
    if not os.path.exists(input_path):
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'

    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Auto-detect client name from content if not provided
    if not client_name:
        match = re.search(r'Client\s*[|:]\s*(.+?)(?:\s*\||$)', md_content)
        if match:
            client_name = match.group(1).strip().strip('[]')

    if not project_name:
        match = re.search(r'Project\s*Name\s*[|:]\s*(.+?)(?:\s*\||$)', md_content)
        if match:
            project_name = match.group(1).strip().strip('[]')

    date_str = datetime.now().strftime('%B %d, %Y')

    doc = Document()
    setup_styles(doc)
    set_page_margins(doc)
    add_cover_page(doc, client_name, project_name, date_str)
    add_table_of_contents(doc)
    parse_and_convert(md_content, doc, client_name, project_name)
    add_footer(doc)

    doc.save(output_path)
    print(f"Document saved: {output_path}")
    return output_path


def merge_parts_to_docx(part_files, output_path, client_name=None, project_name=None):
    """Merge multiple markdown part files into a single professional docx."""
    date_str = datetime.now().strftime('%B %d, %Y')

    # Part titles for dividers
    part_titles = {
        'part1': 'Part 1: Executive Summary, Architecture & ISO Compliance',
        'part2': 'Part 2: Network, Wireless & Server Design',
        'part3': 'Part 3: DR/Backup, Monitoring/SIEM & Power/UPS',
        'part4': 'Part 4: Bill of Materials, Asset Lifecycle & Timeline',
        'part5': 'Part 5: Security Stress Test Results',
        'part6': 'Part 6: Standard Operating Procedures (SOPs)',
    }

    # Sort parts by their part number
    sorted_parts = []
    for f in part_files:
        basename = os.path.basename(f).lower()
        # Extract part number
        match = re.search(r'part\s*(\d+)', basename)
        order = int(match.group(1)) if match else 99
        sorted_parts.append((order, f))
    sorted_parts.sort(key=lambda x: x[0])

    # Auto-detect client/project from first part if not provided
    if sorted_parts and (not client_name or not project_name):
        first_file = sorted_parts[0][1]
        if os.path.exists(first_file):
            with open(first_file, 'r', encoding='utf-8') as f:
                content = f.read()
            if not client_name:
                match = re.search(r'Client\s*[|:]\s*(.+?)(?:\s*\||$)', content)
                if match:
                    client_name = match.group(1).strip().strip('[]')
            if not project_name:
                match = re.search(r'Project\s*Name\s*[|:]\s*(.+?)(?:\s*\||$)', content)
                if match:
                    project_name = match.group(1).strip().strip('[]')

    doc = Document()
    setup_styles(doc)
    set_page_margins(doc)
    add_cover_page(doc, client_name, project_name, date_str)
    add_table_of_contents(doc)

    total_parts = len(sorted_parts)
    missing_parts = [f for _, f in sorted_parts if not os.path.exists(f)]
    if missing_parts:
        print(f"\nWarning: {len(missing_parts)} of {total_parts} part file(s) not found:")
        for f in missing_parts:
            print(f"  - {f}")
        if len(missing_parts) == total_parts:
            print("Error: No part files found. Cannot generate document.")
            sys.exit(1)
        print(f"Continuing with {total_parts - len(missing_parts)} available part(s)...\n")

    added_count = 0
    for idx, (order, filepath) in enumerate(sorted_parts):
        if not os.path.exists(filepath):
            continue

        with open(filepath, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Add part divider
        part_key = f'part{order}'
        title = part_titles.get(part_key, f'Part {order}')
        if idx > 0:
            add_part_divider(doc, title)
        else:
            # First part — just add a subtle header
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = BRAND_ACCENT
            run.font.name = 'Calibri'
            doc.add_paragraph()

        # Parse and add content
        parse_and_convert(md_content, doc, client_name, project_name)

        added_count += 1
        print(f"  Added: {os.path.basename(filepath)} ({title})")

    add_footer(doc)

    doc.save(output_path)
    print(f"\nMerged document saved: {output_path}")
    print(f"  Parts included: {added_count} of {total_parts}")
    if missing_parts:
        print(f"  Parts missing: {len(missing_parts)} (document is INCOMPLETE)")
    print(f"  Client: {client_name or 'Not specified'}")
    print(f"  Project: {project_name or 'Not specified'}")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='VConfi — Convert Markdown Implementation Plan to Word Document'
    )
    subparsers = parser.add_subparsers(dest='command', help='Command to run')

    # Single file conversion (default behavior)
    convert_parser = subparsers.add_parser('convert', help='Convert a single markdown file to docx')
    convert_parser.add_argument('input', help='Input markdown file path')
    convert_parser.add_argument('output', nargs='?', help='Output .docx file path (optional)')
    convert_parser.add_argument('--client', help='Client name for cover page')
    convert_parser.add_argument('--project', help='Project name for cover page')

    # Merge multiple parts
    merge_parser = subparsers.add_parser('merge', help='Merge multiple markdown part files into one docx')
    merge_parser.add_argument('--parts', nargs='+', required=True, help='Part markdown files (in order)')
    merge_parser.add_argument('--output', required=True, help='Output .docx file path')
    merge_parser.add_argument('--client', help='Client name for cover page')
    merge_parser.add_argument('--project', help='Project name for cover page')

    args = parser.parse_args()

    if args.command == 'merge':
        # Expand glob patterns
        import glob
        all_files = []
        for pattern in args.parts:
            matched = glob.glob(pattern)
            if matched:
                all_files.extend(matched)
            else:
                all_files.append(pattern)
        # Remove duplicates while preserving order
        seen = set()
        unique_files = []
        for f in all_files:
            if f not in seen:
                seen.add(f)
                unique_files.append(f)

        print(f"Merging {len(unique_files)} parts into: {args.output}")
        merge_parts_to_docx(unique_files, args.output, args.client, args.project)

    elif args.command == 'convert':
        convert_md_to_docx(args.input, args.output, args.client, args.project)

    else:
        # Default: treat first positional arg as input file (backward compatible)
        # Re-parse without subcommand
        fallback_parser = argparse.ArgumentParser(
            description='VConfi — Convert Markdown Implementation Plan to Word Document'
        )
        fallback_parser.add_argument('input', help='Input markdown file path')
        fallback_parser.add_argument('output', nargs='?', help='Output .docx file path (optional)')
        fallback_parser.add_argument('--client', help='Client name for cover page')
        fallback_parser.add_argument('--project', help='Project name for cover page')

        args = fallback_parser.parse_args()
        convert_md_to_docx(args.input, args.output, args.client, args.project)


if __name__ == '__main__':
    main()
